"""
DOCX document parser using python-docx.
"""

from pathlib import Path
from typing import List, Dict, Any

from ragBaseMaker.parsers.base_parser import BaseParser, ParsedDocument


class DOCXParser(BaseParser):
    """Parser for Microsoft Word documents."""
    
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    
    def parse(self, file_path: str | Path) -> ParsedDocument:
        """
        Parse a DOCX document.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ParsedDocument with extracted text and metadata
        """
        from docx import Document
        
        path = self._validate_file(file_path)
        
        doc = Document(path)
        sections: List[Dict[str, Any]] = []
        all_text_parts: List[str] = []
        
        # Extract core properties as metadata
        metadata: Dict[str, Any] = {}
        if doc.core_properties:
            props = doc.core_properties
            metadata = {
                'author': props.author or '',
                'title': props.title or '',
                'subject': props.subject or '',
                'keywords': props.keywords or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else '',
                'last_modified_by': props.last_modified_by or '',
            }
        
        # Track current section
        current_heading = None
        current_content: List[str] = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if it's a heading
            if para.style and para.style.name.startswith('Heading'):
                # Save previous section if exists
                if current_heading or current_content:
                    sections.append({
                        'type': 'section',
                        'heading': current_heading,
                        'content': '\n'.join(current_content),
                    })
                
                current_heading = text
                current_content = []
                all_text_parts.append(f"\n## {text}\n")
            else:
                current_content.append(text)
                all_text_parts.append(text)
        
        # Save last section
        if current_heading or current_content:
            sections.append({
                'type': 'section',
                'heading': current_heading,
                'content': '\n'.join(current_content),
            })
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables, 1):
            table_data: List[List[str]] = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                sections.append({
                    'type': 'table',
                    'table_number': table_idx,
                    'data': table_data,
                })
                
                # Add table as text
                table_text = '\n'.join([' | '.join(row) for row in table_data])
                all_text_parts.append(f"\n[Table {table_idx}]\n{table_text}\n")
        
        content = self._clean_text('\n'.join(all_text_parts))
        
        return ParsedDocument(
            content=content,
            source_path=str(path.absolute()),
            file_type='docx',
            metadata=metadata,
            title=metadata.get('title') or path.stem,
            sections=sections,
        )

